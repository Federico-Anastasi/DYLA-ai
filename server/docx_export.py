"""Builds the Word version of the brief (brief.json) that gets sent to the client
for validation.

It has to read like a professional document, not like a JSON dump: the reader is a
project manager or a client stakeholder who has never seen the tool. Structure:
title page (carrying a "Draft" notice while the document is not confirmed yet),
table of contents, chapters with their markdown rendered into Word, a Requirements
section, a Glossary, and finally the change history.

DocNotFound and _require are reused from server.exports instead of being duplicated
here: server/main.py already catches `exports.DocNotFound` to turn a missing source
JSON into a 404 for any export, and this module has to keep working with that
mechanism without touching main.py.
"""
from __future__ import annotations

import io
import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .exports import DocNotFound, _require  # noqa: F401  (DocNotFound re-exported)

# ── markdown -> docx (minimal renderer, kept honest by the tests) ──────────

# Only what the chapter bodies of a brief actually need: paragraphs, bulleted and
# numbered lists, inline bold and italic. No headings, links or tables in the body
# markdown — those are modelled as dedicated sections of the document.
_BULLET_RE = re.compile(r"^-\s+")
_NUMBERED_RE = re.compile(r"^\d+\.\s+")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Splits `text` into separate runs so that inline **bold** and *italic* render."""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 3:
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 1:
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def _render_markdown(doc, text: str) -> None:
    """Minimal markdown -> docx renderer: paragraphs (consecutive lines joined by a
    space, separated by a blank line), bulleted lists ('- '), numbered lists ('N. ')
    and inline bold/italic. That covers the chapter bodies of a brief, which are
    prose with the occasional list — this is not a general markdown parser (no
    tables, links, code blocks or nested headings inside the body).
    """
    buffer: list[str] = []

    def flush():
        if buffer:
            p = doc.add_paragraph()
            _add_inline_runs(p, " ".join(buffer))
            buffer.clear()

    for raw_line in (text or "").split("\n"):
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if _BULLET_RE.match(line):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, _BULLET_RE.sub("", line, count=1))
            continue
        m = _NUMBERED_RE.match(line)
        if m:
            flush()
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, line[m.end():])
            continue
        buffer.append(line)
    flush()


# ── Word TOC field ──────────────────────────────────────────────────────────

def _add_toc_field(doc) -> None:
    """Inserts a real TOC \\o "1-3" field (not a static list). python-docx exposes no
    helper for this, so the field XML is assembled by hand (fldChar
    begin/separate/end + instrText). Word fills it in and refreshes it when the
    document is opened (or on F9); the placeholder text below is only what shows
    until it has been updated at least once.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r_element = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Open in Word and refresh the field (F9) to build the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_separate)
    r_element.append(placeholder)
    r_element.append(fld_end)

    # Ask Word to refresh every field on open, so the table of contents fills itself in
    # without the reader needing to know about F9.
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


# ── tables (requirements, glossary, change history) ─────────────────────────

def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"  # built-in docx template style: readable header
    for cell, h in zip(table.rows[0].cells, headers):
        cell.paragraphs[0].add_run(h).bold = True
    for row_vals in rows:
        cells = table.add_row().cells
        for cell, val in zip(cells, row_vals):
            cell.text = "" if val is None else str(val)


# ── title page ──────────────────────────────────────────────────────────────

def _add_titlepage(doc, meta: dict) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(meta["title"])
    title_run.bold = True
    title_run.font.size = Pt(28)

    # A client must never be able to mistake a draft for a validated document, so the
    # notice goes in plain sight, right below the title.
    if meta.get("status") != "confirmed":
        draft_p = doc.add_paragraph()
        draft_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        draft_run = draft_p.add_run("DRAFT — document not validated")
        draft_run.bold = True
        draft_run.font.size = Pt(15)
        draft_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    info_lines = []
    if meta.get("client"):
        info_lines.append(f"Client: {meta['client']}")
    info_lines.append(f"Project: {meta['project']}")
    info_lines.append(f"Date: {meta['date']}")
    if meta.get("version") is not None:
        info_lines.append(f"Version: {meta['version']}")
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    if meta.get("notes"):
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_p.add_run(meta["notes"]).italic = True

    doc.add_page_break()


# ── entry point ──────────────────────────────────────────────────────────

def brief_docx(project: str) -> bytes:
    data = _require(project, "brief")
    meta = data["meta"]
    chapters = data.get("chapters") or []
    requirements = data.get("requirements") or []
    glossary = data.get("glossary") or []
    changelog = data.get("changelog") or []

    doc = docx.Document()

    _add_titlepage(doc, meta)

    # Table of contents: a "title-looking" paragraph (not a real heading, otherwise the
    # TOC field below would list itself) followed by Word's actual TOC field.
    idx_p = doc.add_paragraph()
    idx_run = idx_p.add_run("Table of Contents")
    idx_run.bold = True
    idx_run.font.size = Pt(20)
    _add_toc_field(doc)
    doc.add_page_break()

    for chapter in chapters:
        doc.add_heading(chapter["title"], level=chapter.get("level", 1))
        if chapter.get("open"):
            note_p = doc.add_paragraph()
            note_run = note_p.add_run(
                "Incomplete chapter: there are still open questions to clear with the client."
            )
            note_run.italic = True
            note_run.font.color.rgb = RGBColor(0xB8, 0x86, 0x00)
        _render_markdown(doc, chapter.get("body", ""))

    if requirements:
        doc.add_heading("Requirements", level=1)
        _add_table(
            doc,
            ["ID", "Title", "Description", "Priority", "Status"],
            [
                [r["id"], r["title"], r["description"], r.get("priority") or "",
                 r.get("status") or ""]
                for r in requirements
            ],
        )

    if glossary:
        doc.add_heading("Glossary", level=1)
        _add_table(
            doc,
            ["Term", "Definition"],
            [[g["term"], g["definition"]] for g in glossary],
        )

    if changelog:
        # A real page footer repeats identically on every page, which does not suit a
        # list that grows with every meeting folded in. The history works better as a
        # standalone closing section.
        doc.add_heading("Change history", level=1)
        _add_table(
            doc,
            ["Date", "Source", "What changed"],
            [[c["date"], c["source"], c["summary"]] for c in changelog],
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
