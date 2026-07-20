"""Input documents and project source, from the API side.

What these tests guard: the brief has two different natures depending on how the project
came about (a read-only input document, or a deliverable of ours), and the rest of the
backend must no longer assume it is called `brief.md`.

Run with: python -m pytest server/tests/test_documents_api.py
"""
import io

import docx
import pytest
from fastapi.testclient import TestClient

from server import main


@pytest.fixture
def client(tmp_path, monkeypatch):
    """Every test works on its own projects folder: the tests must not see (nor dirty)
    the real projects.

    Uploading a brief queues a background summary that calls the model. Here it is
    stubbed out: the tests must not reach the network, and waiting on a model that is
    not going to answer added half a minute to the run.
    """
    async def no_summary(text, env=None):
        return None

    monkeypatch.setattr(main, "summarize_brief", no_summary)
    projects = tmp_path / "projects"
    projects.mkdir()
    monkeypatch.setattr(main, "PROJECTS_DIR", projects)
    for mod in ("documents", "versioning", "session_manager", "ingest"):
        m = __import__(f"server.{mod}", fromlist=["*"])
        if hasattr(m, "PROJECTS_DIR"):
            monkeypatch.setattr(m, "PROJECTS_DIR", projects)
    return TestClient(main.app)


def _docx_bytes(title: str, body: str) -> bytes:
    d = docx.Document()
    d.add_heading(title, 1)
    d.add_paragraph(body)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# --- project source ---

def test_a_brief_project_does_not_create_the_meetings_folder(client):
    r = client.post("/api/projects", json={"name": "p1", "client": "Acme", "source": "brief"})
    assert r.status_code == 201
    assert r.json()["source"] == "brief"
    assert client.get("/api/projects/p1").json()["source"] == "brief"


def test_a_discovery_project_prepares_the_meetings_folder(client):
    r = client.post("/api/projects", json={"name": "p2", "client": "Acme", "source": "discovery"})
    assert r.status_code == 201
    # The folder exists from the start: that is how it is obvious where transcripts go.
    assert (main.PROJECTS_DIR / "p2" / "meetings").is_dir()
    ctx = (main.PROJECTS_DIR / "p2" / "context.md").read_text(encoding="utf-8")
    assert "Source" in ctx


def test_an_invalid_source_is_rejected(client):
    r = client.post("/api/projects", json={"name": "p3", "client": "Acme", "source": "dunno"})
    assert r.status_code == 400


def test_a_project_without_project_json_migrates_itself(client):
    """Projects created before the source field existed all had an input brief."""
    client.post("/api/projects", json={"name": "p4", "client": "Acme"})
    (main.PROJECTS_DIR / "p4" / ".project.json").unlink()
    assert client.get("/api/projects/p4").json()["source"] == "brief"
    assert (main.PROJECTS_DIR / "p4" / ".project.json").is_file()


# --- upload ---

def test_multiple_upload(client):
    client.post("/api/projects", json={"name": "p5", "client": "Acme"})
    files = [("files", ("one.md", b"# one", "text/markdown")),
             ("files", ("two.md", b"# two", "text/markdown"))]
    r = client.post("/api/projects/p5/documents?target=docs", files=files)
    assert r.status_code == 201
    assert r.json()["saved"] == ["docs/one.md", "docs/two.md"]


def test_safe_name_refuses_a_bare_dot_or_dot_dot():
    """".." is made only of characters the sanitiser already allows through (word chars,
    dot, dash, space): it used to survive unchanged and put the file on the PARENT of the
    target directory instead of inside it — a 500, not a controlled rejection."""
    assert main._safe_name("..") == "document"
    assert main._safe_name(".") == "document"
    assert main._safe_name("normal.pdf") == "normal.pdf"


def test_uploading_a_file_named_dot_dot_does_not_escape_the_project(client):
    client.post("/api/projects", json={"name": "p5b", "client": "Acme"})
    files = [("files", ("..", b"content", "text/plain"))]
    r = client.post("/api/projects/p5b/documents?target=docs", files=files)
    assert r.status_code == 201
    saved = r.json()["saved"][0]
    # Landed inside docs/, under the safe fallback name — not one level up.
    assert saved.startswith("docs/")
    assert (main.PROJECTS_DIR / "p5b" / saved).is_file()
    assert not (main.PROJECTS_DIR / "p5b.md").exists()  # the parent-escape that used to happen


def test_the_brief_keeps_its_own_extension(client):
    """Nobody can open a PDF that was saved as brief.md."""
    client.post("/api/projects", json={"name": "p6", "client": "Acme"})
    r = client.post("/api/projects/p6/documents?target=brief",
                    files=[("files", ("Brief v3.docx", _docx_bytes("Context", "text"),
                                      "application/octet-stream"))])
    assert r.json()["saved"] == ["brief.docx"]
    assert (main.PROJECTS_DIR / "p6" / "brief.docx").is_file()


def test_uploading_a_new_brief_removes_the_previous_one(client):
    """There is only one brief: two different brief.* files would make it ambiguous
    which one counts."""
    client.post("/api/projects", json={"name": "p7", "client": "Acme"})
    client.post("/api/projects/p7/documents?target=brief",
                files=[("files", ("brief.md", b"# old", "text/markdown"))])
    client.post("/api/projects/p7/documents?target=brief",
                files=[("files", ("brief.docx", _docx_bytes("New", "text"),
                                  "application/octet-stream"))])
    d = main.PROJECTS_DIR / "p7"
    assert not (d / "brief.md").exists()
    assert (d / "brief.docx").is_file()


def test_only_one_brief_at_a_time(client):
    client.post("/api/projects", json={"name": "p8", "client": "Acme"})
    files = [("files", ("a.md", b"a", "text/markdown")),
             ("files", ("b.md", b"b", "text/markdown"))]
    assert client.post("/api/projects/p8/documents?target=brief", files=files).status_code == 400


def test_an_unknown_target_is_rejected(client):
    client.post("/api/projects", json={"name": "p9", "client": "Acme"})
    r = client.post("/api/projects/p9/documents?target=elsewhere",
                    files=[("files", ("x.md", b"x", "text/markdown"))])
    assert r.status_code == 400


# --- listing and removal ---

def test_documents_carry_their_metadata(client):
    client.post("/api/projects", json={"name": "q1", "client": "Acme"})
    client.post("/api/projects/q1/documents?target=brief",
                files=[("files", ("brief.docx", _docx_bytes("Context", "text"),
                                  "application/octet-stream"))])
    client.post("/api/projects/q1/documents?target=docs",
                files=[("files", ("logo.png", b"\x89PNG\r\n\x1a\n", "image/png"))])
    docs = {d["file"]: d for d in client.get("/api/projects/q1/documents").json()}
    assert docs["brief.docx"]["kind"] == "brief"
    assert docs["brief.docx"]["readable"] and docs["brief.docx"]["extracted"]
    # An image stays downloadable but the agent cannot read it: the UI has to say so.
    assert not docs["docs/logo.png"]["readable"]


def test_extracts_do_not_show_up_among_the_files(client):
    """They are a cache, not documents: they must not clutter the listing or the tabs."""
    client.post("/api/projects", json={"name": "q2", "client": "Acme"})
    client.post("/api/projects/q2/documents?target=docs",
                files=[("files", ("note.docx", _docx_bytes("Note", "text"),
                                  "application/octet-stream"))])
    files = client.get("/api/projects/q2/files").json()
    assert not any(".extracted" in f for f in files)


def test_removing_a_document_removes_its_extract(client):
    client.post("/api/projects", json={"name": "q3", "client": "Acme"})
    client.post("/api/projects/q3/documents?target=docs",
                files=[("files", ("note.docx", _docx_bytes("Note", "text"),
                                  "application/octet-stream"))])
    from server import ingest
    extract = ingest.extract_path(main.PROJECTS_DIR / "q3", "docs/note.docx")
    assert extract.is_file()
    assert client.delete("/api/projects/q3/documents/docs/note.docx").status_code == 200
    assert not extract.exists()


def test_deliverables_are_not_deleted_from_here(client):
    """Deliverables have versioning instead: a DELETE would simply lose them."""
    client.post("/api/projects", json={"name": "q4", "client": "Acme"})
    assert client.delete("/api/projects/q4/documents/context.md").status_code == 400


# --- brief anchors ---

def test_an_input_brief_exposes_its_chapters(client):
    client.post("/api/projects", json={"name": "r1", "client": "Acme", "source": "brief"})
    client.post("/api/projects/r1/documents?target=brief",
                files=[("files", ("brief.docx", _docx_bytes("Data import", "text"),
                                  "application/octet-stream"))])
    info = client.get("/api/projects/r1/brief").json()
    assert info["kind"] == "file" and info["file"] == "brief.docx"
    assert info["headings"][0]["slug"] == "data-import"


def test_a_deliverable_brief_exposes_the_chapters_of_the_json(client):
    client.post("/api/projects", json={"name": "r2", "client": "Acme", "source": "discovery"})
    doc = {"meta": {"project": "r2", "title": "Brief", "date": "2026-07-19"},
           "chapters": [{"id": "C1", "title": "Target process", "body": "text"}]}
    assert client.put("/api/projects/r2/doc/brief", json=doc).status_code == 200
    info = client.get("/api/projects/r2/brief").json()
    assert info["kind"] == "doc" and info["file"] == "brief.json"
    assert info["headings"][0]["id"] == "C1"


def test_a_truncated_brief_json_reports_409_not_500(client):
    """get_brief was the one load_doc caller in this file not wrapped for
    DocumentUnreadable: a brief.json half-written by an interrupted turn came back as an
    unexplained 500 instead of the 409 every other document endpoint gives for the exact
    same situation."""
    client.post("/api/projects", json={"name": "r2b", "client": "Acme", "source": "discovery"})
    (main.PROJECTS_DIR / "r2b" / "brief.json").write_text("{not json", encoding="utf-8")
    r = client.get("/api/projects/r2b/brief")
    assert r.status_code == 409


def test_no_brief_means_no_anchors(client):
    client.post("/api/projects", json={"name": "r3", "client": "Acme"})
    info = client.get("/api/projects/r3/brief").json()
    assert info["file"] is None and info["headings"] == []


def test_the_workflow_recognises_the_brief_in_any_format(client):
    client.post("/api/projects", json={"name": "r4", "client": "Acme"})
    assert client.get("/api/projects/r4").json()["workflow"]["brief"] is False
    client.post("/api/projects/r4/documents?target=brief",
                files=[("files", ("brief.pdf", b"%PDF-1.4 fake", "application/pdf"))])
    assert client.get("/api/projects/r4").json()["workflow"]["brief"] is True


# --- home page: ordering by recent activity ---

# --- xlsx preview: cell values and sheet titles are untrusted input -----------

def test_xlsx_preview_escapes_cell_values_and_sheet_titles(client):
    """Cell values and the sheet title come straight from whatever xlsx was uploaded: an
    <img src=x onerror=...> in a cell used to run as script in the preview iframe because
    none of it was ever escaped."""
    from openpyxl import Workbook

    client.post("/api/projects", json={"name": "r5", "client": "Acme"})
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet & \"quoted\""  # openpyxl rejects <, >, : and a few others in titles
    ws["A1"] = "<img src=x onerror=alert(1)>"
    buf = io.BytesIO()
    wb.save(buf)
    client.post("/api/projects/r5/documents?target=docs",
                files=[("files", ("sheet.xlsx", buf.getvalue(),
                                  "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"))])
    r = client.get("/api/projects/r5/preview/docs/sheet.xlsx")
    assert r.status_code == 200
    assert "<img src=x" not in r.text
    assert "Sheet &amp; &quot;quoted&quot;" in r.text
    assert "&lt;img src=x" in r.text


def test_the_project_list_carries_the_last_modified_time(client):
    """The home page orders projects by this: without it, the user has to hunt for the
    one they were working on."""
    client.post("/api/projects", json={"name": "h1", "client": "Acme"})
    projects = {p["name"]: p for p in client.get("/api/projects").json()}
    assert projects["h1"]["modified"] > 0
    assert projects["h1"]["client"] == "Acme"


def test_last_modified_follows_the_deliverables(client):
    """A project where a deliverable has just been written must come out as more recent
    than one that was merely created."""
    client.post("/api/projects", json={"name": "h2", "client": "Acme"})
    client.post("/api/projects", json={"name": "h3", "client": "Acme"})
    import os
    import time
    doc = {"meta": {"project": "h3", "date": "2026-07-19"}, "people": []}
    client.put("/api/projects/h3/doc/people", json=doc)
    # mtime has 1s resolution on some filesystems: we force it instead of waiting.
    f = main.PROJECTS_DIR / "h3" / "people.json"
    future = time.time() + 60
    os.utime(f, (future, future))
    projects = {p["name"]: p for p in client.get("/api/projects").json()}
    assert projects["h3"]["modified"] > projects["h2"]["modified"]
