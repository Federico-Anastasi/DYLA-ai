"""Tests for the Word export of the brief (server/docx_export.py) and the test plan
xlsx (at the end of server/exports.py), plus a regression on the timeline summary.

Fixture pattern taken from test_mockup_export.py: it is `exports.load_doc` that gets
monkeypatched, not the calling modules, because both `docx_export._require` and
`exports._require` resolve the name `load_doc` in the global namespace of
`server/exports.py` — patching it there covers both modules.
"""
import io
import xml.etree.ElementTree as ET

import docx
import openpyxl
import pytest

from server import docx_export, exports


@pytest.fixture
def patch_docs(monkeypatch):
    """Makes load_doc('t', <name>) resolve to the document in `docs[<name>]`, None
    otherwise."""

    def _patch(docs: dict):
        def _load(project, name):
            if project != "t":
                return None
            return docs.get(name)
        monkeypatch.setattr(exports, "load_doc", _load)

    return _patch


# --- minimal document fixtures ------------------------------------------------

def _brief_doc(status: str = "draft") -> dict:
    return {
        "meta": {
            "project": "t", "title": "Booking Portal Brief", "date": "2026-07-19",
            "client": "Acme", "status": status, "version": 2,
            "notes": "Out of scope: booking a meeting across two rooms.",
        },
        "chapters": [
            {
                "id": "C1", "title": "Context and goals", "level": 1,
                "body": (
                    "The project covers **the booking** of meeting rooms.\n\n"
                    "Key points:\n- First point\n- Second point\n\n"
                    "Some steps *in italics* for detail."
                ),
                "sources": ["meeting1.md"],
            },
            {
                "id": "C2", "title": "Open flow", "level": 1,
                "body": "1. Step one\n2. Step two",
                "open": True,
            },
        ],
        "requirements": [
            {
                "id": "R1", "title": "Change history",
                "description": "The system must record every state change with the user and the date.",
                "chapter": "C1", "priority": "high", "status": "validated",
            },
        ],
        "glossary": [{"term": "Slot", "definition": "A bookable time window on one room"}],
        "changelog": [{"date": "2026-07-19", "source": "meeting1.md", "summary": "First draft"}],
    }


def _test_plan_doc() -> dict:
    return {
        "meta": {"project": "t", "title": "Booking Portal test plan",
                 "date": "2026-07-19", "client": "Acme"},
        "cases": [
            {
                "id": "TC1", "title": "Valid booking created", "epic": "E1", "task": "E1.T1",
                "type": "functional", "preconditions": "Signed-in user with the Employee role",
                "steps": [
                    {"n": 1, "action": "Open the New Booking page",
                     "expected": "The form opens empty"},
                    {"n": 2, "action": "Fill the required fields and click Save",
                     "expected": "The system shows the save confirmation"},
                ],
                "expected_result": "The booking exists and shows up in the calendar",
                "outcome": "ok", "tester": "MR", "run_at": "2026-07-18",
            },
            {
                "id": "TC2", "title": "Overlapping booking blocked", "epic": "E1",
                "type": "negative", "preconditions": "A booking already exists on that slot",
                "steps": [
                    {"n": 1, "action": "Book the same room on an overlapping slot",
                     "expected": "The system blocks the save"},
                ],
                "expected_result": "Only the original booking survives",
                "outcome": "ko", "tester": "MR", "run_at": "2026-07-18",
            },
        ],
    }


def _estimate_doc() -> dict:
    return {
        "meta": {"project": "t", "title": "Estimate", "date": "2026-07-19",
                 "contingency_pct": 15},
        "epics": [
            {
                "id": "E1", "name": "1. Booking",
                "tasks": [
                    {"id": "E1.T1", "task": "Booking form", "days": 1.5,
                     "description": "", "dev_tasks": []},
                ],
                "e2e": {"label": "E2E test of the Booking flow", "days": 0.5},
            },
            {
                "id": "E2", "name": "2. Notifications",
                "tasks": [
                    {"id": "E2.T1", "task": "Reminder notification", "days": 1.0,
                     "description": "", "dev_tasks": []},
                ],
                "e2e": {"label": "E2E test of the Notifications flow", "days": 0.5},
            },
        ],
    }


# --- brief docx ---------------------------------------------------------------

def test_the_docx_opens_with_its_headings_and_markdown(patch_docs):
    patch_docs({"brief": _brief_doc()})
    content = docx_export.brief_docx("t")
    document = docx.Document(io.BytesIO(content))

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Context and goals" in full_text
    assert "Open flow" in full_text
    assert "incomplete" in full_text.lower()  # the note on the open chapter

    bullet_paragraphs = [p for p in document.paragraphs if p.style.name == "List Bullet"]
    assert any("First point" in p.text for p in bullet_paragraphs)
    assert any("Second point" in p.text for p in bullet_paragraphs)

    numbered_paragraphs = [p for p in document.paragraphs if p.style.name == "List Number"]
    assert any("Step one" in p.text for p in numbered_paragraphs)
    assert any("Step two" in p.text for p in numbered_paragraphs)

    bold_runs = [r for p in document.paragraphs for r in p.runs if r.bold]
    assert any("the booking" in r.text.lower() for r in bold_runs)

    italic_runs = [r for p in document.paragraphs for r in p.runs if r.italic]
    assert any("in italics" in r.text.lower() for r in italic_runs)

    # Requirements + Glossary + Change history = 3 tables
    assert len(document.tables) == 3
    req_table = document.tables[0]
    assert req_table.rows[1].cells[0].text == "R1"


def test_the_title_page_carries_the_notes_from_meta(patch_docs):
    """meta.notes is what says what the document deliberately leaves out. It used to be
    read under the wrong key, so it silently never reached the page."""
    patch_docs({"brief": _brief_doc()})
    content = docx_export.brief_docx("t")
    text = "\n".join(p.text for p in docx.Document(io.BytesIO(content)).paragraphs)
    assert "Out of scope: booking a meeting across two rooms." in text


def test_the_draft_notice_is_there_when_not_confirmed(patch_docs):
    patch_docs({"brief": _brief_doc(status="draft")})
    content = docx_export.brief_docx("t")
    text = "\n".join(p.text for p in docx.Document(io.BytesIO(content)).paragraphs)
    assert "DRAFT" in text


def test_the_draft_notice_is_gone_once_confirmed(patch_docs):
    patch_docs({"brief": _brief_doc(status="confirmed")})
    content = docx_export.brief_docx("t")
    text = "\n".join(p.text for p in docx.Document(io.BytesIO(content)).paragraphs)
    assert "DRAFT" not in text


def test_brief_docx_doc_not_found():
    with pytest.raises(docx_export.DocNotFound):
        docx_export.brief_docx("nonexistent_project_xyz_pytest")


# --- test plan xlsx -----------------------------------------------------------

def test_the_test_plan_has_one_row_per_step(patch_docs):
    patch_docs({"test_plan": _test_plan_doc()})
    content = exports.test_plan_xlsx("t")
    wb = openpyxl.load_workbook(io.BytesIO(content))

    assert wb.sheetnames[0] == "Test plan"
    ws = wb["Test plan"]
    # TC1 has 2 steps, TC2 has 1: 3 data rows plus the header
    assert ws.max_row == 4
    assert ws.cell(row=1, column=1).value == "Case ID"
    assert ws.cell(row=2, column=1).value == "TC1"
    assert ws.cell(row=2, column=7).value == 1  # step no.
    assert ws.cell(row=3, column=7).value == 2
    assert ws.cell(row=4, column=1).value == "TC2"

    # Outcome colouring: green for TC1 (ok), red for TC2 (ko)
    assert ws.cell(row=2, column=11).fill.fgColor.rgb.endswith("C6EFCE")
    assert ws.cell(row=4, column=11).fill.fgColor.rgb.endswith("FFC7CE")


def test_the_coverage_sheet_flags_an_uncovered_epic(patch_docs):
    patch_docs({"test_plan": _test_plan_doc(), "estimate": _estimate_doc()})
    content = exports.test_plan_xlsx("t")
    wb = openpyxl.load_workbook(io.BytesIO(content))

    assert "Coverage" in wb.sheetnames
    ws2 = wb["Coverage"]
    rows = {row[0].value: row for row in ws2.iter_rows(min_row=2)}

    e1 = rows["E1"]
    assert e1[2].value == 2  # no. of cases
    assert e1[3].value == 1  # OK
    assert e1[4].value == 1  # KO

    e2 = rows["E2"]
    assert e2[2].value == 0  # no cases at all: the epic is uncovered
    assert e2[0].fill.fgColor.rgb.endswith("FFC7CE")


def test_without_an_estimate_the_coverage_sheet_is_skipped(patch_docs):
    patch_docs({"test_plan": _test_plan_doc()})  # no 'estimate'
    content = exports.test_plan_xlsx("t")
    wb = openpyxl.load_workbook(io.BytesIO(content))
    assert "Coverage" not in wb.sheetnames


def test_test_plan_xlsx_doc_not_found():
    with pytest.raises(exports.DocNotFound):
        exports.test_plan_xlsx("nonexistent_project_xyz_pytest")


# --- regression: the timeline summary -----------------------------------------

def test_the_timeline_summary_reports_the_duration(patch_docs):
    """The Summary sheet used to compute the duration from an attribute Plan does not
    have: the export died with a 500 as soon as anyone downloaded the timeline of a
    planned project."""
    estimate = {
        "meta": {"project": "t", "title": "E", "date": "2026-07-19", "contingency_pct": 15},
        "epics": [{
            "id": "E1", "name": "1. Epic",
            "tasks": [{"id": "E1.T1", "task": "Task", "days": 2, "description": "d",
                       "dev_tasks": [{"id": "E1.T1.D1", "dev_task": "Dev", "description": "d",
                                      "days": 2, "layer": 1}]}],
            "e2e": {"label": "E2E test", "days": 1},
        }],
    }
    timeline = {
        "meta": {"project": "t", "date": "2026-07-19"},
        "start_date": "2026-09-07",
        "team": [{"id": "dev1", "name": "Anna"}],
        "lanes": [{"dev": "dev1", "items": ["E1.T1.D1", "E1.E2E"]}],
    }
    patch_docs({"estimate": estimate, "timeline": timeline})
    wb = openpyxl.load_workbook(io.BytesIO(exports.timeline_xlsx("t")))
    summary = wb["Summary"]
    labels = {row[0]: row[1] for row in summary.iter_rows(values_only=True) if row[0]}
    # Both ends included: 3 days on one person starting on a Monday = Monday to Wednesday.
    assert labels["Duration (calendar days)"] == 3


# --- drawio: names come from the document, and documents contain punctuation ---------

def _data_model(area_name="Core", table_name="Client", field_name="email"):
    return {
        "meta": {"title": "T", "date": "2026-07-20", "status": "draft"},
        "areas": [{"id": "a1", "name": area_name, "color": "#4472C4"}],
        "tables": [{"id": "client", "name": table_name, "area": "a1",
                    "fields": [{"name": field_name, "type": "TEXT"}]}],
        "relations": [],
    }


def test_an_ampersand_in_a_name_does_not_break_the_diagram(patch_docs):
    """Found on a real document: an area called "Stock & Inventory" — an entirely ordinary
    name — produced a .drawio that would not parse, and the export answered 500 with no
    explanation. Every label in the diagram comes from the document, so every label has to
    be escaped."""
    patch_docs({"data_model": _data_model(area_name="Stock & Inventory")})
    xml = exports.data_model_drawio("t")
    ET.fromstring(xml)  # would raise if the file were malformed
    assert "Stock &amp; Inventory" in xml


def test_angle_brackets_and_quotes_survive_too(patch_docs):
    """The same hole, reached by other characters: a type like List<String>, or a name
    with quotes in it."""
    patch_docs({"data_model": _data_model(table_name='Order "line"', field_name="List<String>")})
    xml = exports.data_model_drawio("t")
    ET.fromstring(xml)
    assert "&lt;String&gt;" in xml
    assert "&quot;line&quot;" in xml


def test_the_diagram_still_says_what_it_should(patch_docs):
    """Escaping must not turn into mangling: the words are still readable."""
    patch_docs({"data_model": _data_model()})
    xml = exports.data_model_drawio("t")
    root = ET.fromstring(xml)
    values = {c.get("value") for c in root.iter("mxCell")}
    assert "Client" in values
    assert "email" in values


def test_a_relation_with_no_dot_does_not_break_the_diagram(patch_docs):
    """relations[].from/to is a plain string in the schema, no pattern enforcing
    "table.field": a relation written as "Orders" instead of "Orders.id" used to raise
    ValueError from the tuple-unpacking split() and take the whole export down with it."""
    data = _data_model()
    data["relations"] = [{"from": "client", "to": "client.id", "type": "1-n"}]
    patch_docs({"data_model": data})
    xml = exports.data_model_drawio("t")
    ET.fromstring(xml)  # would not even get here before the fix


def test_field_notes_reach_the_data_model_sheet(patch_docs):
    """The schema calls the field "notes" (plural); reading "note" always returned None,
    so the Notes column of the Data Model sheet was silently empty."""
    data = _estimate_doc()
    dm = _data_model()
    dm["tables"][0]["fields"][0]["notes"] = "Unique per tenant"
    patch_docs({"estimate": data, "data_model": dm})
    wb = openpyxl.load_workbook(io.BytesIO(exports.estimate_xlsx("t")))
    ws2 = wb["Data Model"]
    notes_col = [row[6] for row in ws2.iter_rows(min_row=3, values_only=True)]
    assert "Unique per tenant" in notes_col


# --- an epic with no tasks must not swallow the next epic's name --------------

def _estimate_with_an_empty_epic() -> dict:
    return {
        "meta": {"project": "t", "title": "Estimate", "date": "2026-07-19",
                 "contingency_pct": 15},
        "epics": [
            {"id": "E1", "name": "1. Empty epic", "tasks": []},
            {
                "id": "E2", "name": "2. Notifications",
                "tasks": [{"id": "E2.T1", "task": "Reminder", "days": 1.0,
                           "description": "", "dev_tasks": []}],
                "e2e": {"label": "E2E test of the Notifications flow", "days": 0.5},
            },
        ],
    }


def test_an_empty_epic_does_not_erase_the_next_ones_name(patch_docs):
    """Neither 'tasks' nor 'e2e' is required to be non-empty by the schema. An epic with
    no rows of its own used to leave start > end: the merge was skipped and the epic's
    name landed on the row the NEXT epic was about to claim, overwriting it — the empty
    epic vanished from the file handed to the client."""
    patch_docs({"estimate": _estimate_with_an_empty_epic()})
    wb = openpyxl.load_workbook(io.BytesIO(exports.estimate_xlsx("t")))
    ws = wb.active
    names_in_col_a = {ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)}
    assert "1. Empty epic" in names_in_col_a
    assert "2. Notifications" in names_in_col_a


def test_dev_tasks_xlsx_survives_an_epic_with_no_tasks(patch_docs):
    patch_docs({"estimate": _estimate_with_an_empty_epic()})
    wb = openpyxl.load_workbook(io.BytesIO(exports.dev_tasks_xlsx("t")))
    ws = wb.active
    names_in_col_a = {ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)}
    assert "1. Empty epic" in names_in_col_a
    assert "2. Notifications" in names_in_col_a


def test_dev_tasks_xlsx_auto_filter_matches_its_six_columns(patch_docs):
    """The sheet has 6 columns (A-F); the filter used to be set to A1:G, one column past
    the last real one."""
    patch_docs({"estimate": _estimate_doc()})
    wb = openpyxl.load_workbook(io.BytesIO(exports.dev_tasks_xlsx("t")))
    ws = wb.active
    assert ws.auto_filter.ref.startswith("A1:F")
