"""Text extraction from input documents, and the anchors used for citations.

The case that really matters is the scanned PDF: without the empty-text check the
extract would be nothing but a run of page markers, and the interface would call
"readable" a document in which the agent cannot find a single word.

Run with: python -m pytest server/tests/test_ingest.py
"""
from pathlib import Path

import docx
import openpyxl
import pytest
from pypdf import PdfWriter

from server import ingest


@pytest.fixture
def project(tmp_path):
    d = tmp_path / "proj"
    (d / "docs").mkdir(parents=True)
    return d


def _docx_with_chapters(path):
    doc = docx.Document()
    doc.add_heading("Project context", 1)
    doc.add_paragraph("The team currently tracks bookings in a spreadsheet.")
    doc.add_heading("Data import", 2)
    doc.add_paragraph("The file arrives from the facilities system.")
    doc.save(str(path))


# --- extraction ---

def test_a_docx_becomes_markdown_with_headings(project):
    _docx_with_chapters(project / "brief.docx")
    text = ingest.extract_text(project / "brief.docx")
    assert "# Project context" in text
    assert "## Data import" in text
    assert "facilities system" in text


def test_a_pdf_with_no_text_produces_no_extract(project):
    """A scanned PDF: pages that are images, no extractable text."""
    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    with (project / "docs" / "scan.pdf").open("wb") as fh:
        w.write(fh)
    assert ingest.extract_text(project / "docs" / "scan.pdf") is None
    assert ingest.build_extract(project, "docs/scan.pdf") is None


def test_an_unsupported_format_stays_without_an_extract(project):
    (project / "docs" / "image.png").write_bytes(b"\x89PNG\r\n\x1a\n")
    assert ingest.extract_text(project / "docs" / "image.png") is None


def test_a_corrupt_file_does_not_raise(project):
    """An unreadable document must not fail the upload: it just ends up without an
    extract."""
    (project / "docs" / "broken.pdf").write_bytes(b"i am not a pdf")
    assert ingest.extract_text(project / "docs" / "broken.pdf") is None


def test_from_xlsx_closes_the_workbook_even_when_reading_a_sheet_fails(monkeypatch):
    """read_only mode needs an explicit close() to release its file handle; that call
    used to sit only on the happy path, so a workbook that raised partway through a
    sheet stayed open for the rest of the process — on Windows, long enough that the
    document could no longer be deleted or replaced."""
    closed = {"called": False}

    class _FakeSheet:
        title = "Sheet1"

        def iter_rows(self, values_only=True):
            raise RuntimeError("corrupt sheet")

    class _FakeWorkbook:
        worksheets = [_FakeSheet()]

        def close(self):
            closed["called"] = True

    monkeypatch.setattr(openpyxl, "load_workbook", lambda *a, **k: _FakeWorkbook())

    with pytest.raises(RuntimeError):
        ingest._from_xlsx(Path("whatever.xlsx"))
    assert closed["called"] is True


def test_the_extract_on_disk_and_its_removal(project):
    _docx_with_chapters(project / "docs" / "flows.docx")
    dest = ingest.build_extract(project, "docs/flows.docx")
    assert dest is not None and dest.is_file()
    # Flat path under .extracted/, so no parallel trees get created.
    assert dest.name == "docs__flows.docx.md"
    assert "automatic extract" in dest.read_text(encoding="utf-8")
    ingest.drop_extract(project, "docs/flows.docx")
    assert not dest.exists()


def test_readable_text_builds_the_extract_if_it_is_missing(project):
    _docx_with_chapters(project / "docs" / "minutes.docx")
    text = ingest.readable_text(project, "docs/minutes.docx")
    assert "Project context" in text
    assert ingest.extract_path(project, "docs/minutes.docx").is_file()


def test_readable_text_reads_plain_text_directly(project):
    (project / "docs" / "note.md").write_text("# Note\n\ncontent", encoding="utf-8")
    assert ingest.readable_text(project, "docs/note.md") == "# Note\n\ncontent"
    # A file that is already text needs no extract.
    assert not ingest.extract_path(project, "docs/note.md").exists()


# --- the brief can come in several formats ---

def test_brief_file_recognises_the_format(project):
    assert ingest.brief_file(project) is None
    (project / "brief.pdf").write_bytes(b"%PDF-1.4")
    assert ingest.brief_file(project) == "brief.pdf"


def test_markdown_wins_over_the_binary(project):
    """If there were two brief.* files, the one we write ourselves takes precedence."""
    (project / "brief.pdf").write_bytes(b"%PDF-1.4")
    (project / "brief.md").write_text("# Brief", encoding="utf-8")
    assert ingest.brief_file(project) == "brief.md"


# --- anchors ---

def test_headings_carry_slug_line_and_page():
    text = "<!-- page 1 -->\n# First\ntext\n\n<!-- page 4 -->\n## Second chapter\n"
    h = ingest.headings(text)
    assert [x["title"] for x in h] == ["First", "Second chapter"]
    assert h[1]["slug"] == "second-chapter"
    assert h[1]["level"] == 2
    # The page is what PDFs need: the navigable anchor there is #page=N.
    assert h[0]["page"] == 1 and h[1]["page"] == 4


def test_duplicate_headings_keep_distinct_slugs():
    """Two chapters with the same title must get different anchors, otherwise the
    citation always lands on the first one."""
    h = ingest.headings("# Foreword\n\n# Foreword\n")
    assert [x["slug"] for x in h] == ["foreword", "foreword-1"]


def test_slugify_handles_accents_and_punctuation():
    assert ingest.slugify("Activité de l'équipe: 3rd level") == "activité-de-léquipe-3rd-level"


def test_headings_without_pages_invent_no_numbers():
    h = ingest.headings("# Title\n")
    assert h[0]["page"] is None and h[0]["line"] == 1
